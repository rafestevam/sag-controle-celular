from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import app_config
import os
from datetime import datetime

class DocumentComposer:

    @classmethod
    def cria_documento(cls, funcionario):
        #Criação do Documento
        documento = Document()

        # Definição de Margem do Documento
        secao_corrente = documento.sections[-1]
        secao_corrente.top_margin = Inches(0.5)
        secao_corrente.bottom_margin = Inches(0.5)
        secao_corrente.left_margin = Inches(0.5)
        secao_corrente.right_margin = Inches(0.5)

        # Formatações gerais do Documento
        documento.add_picture(f'{app_config.IMG_LOCATION}/sag-dark-logo.png', width=Inches(1.71))
        estilo = documento.styles['Normal']
        fonte = estilo.font
        fonte.name = 'Arial Narrow'
        fonte.size = Pt(9)
        
        # Titulo do Documento
        titulo = documento.add_paragraph()
        titulo_format = titulo.paragraph_format
        titulo_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        titulo_run = titulo.add_run('TERMO DE RESPONSABILIDADE PARA USO DE DISPOSITIVO MÓVEL')
        titulo_run.bold = True
        titulo_font = titulo_run.font
        titulo_font.name = 'Arial Narrow'
        titulo_font.size = Pt(10)

        par_1 = documento.add_paragraph()
        par_1.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        run_1 = 'Pelo presente Termo de Responsabilidade para uso de dispositivo móvel, ' \
                'as partes, de um lado, '
        run_2 = f'{funcionario.nome} {funcionario.sobrenome}, '
        run_3 = 'funcionário, brasileiro, portador do RG nº '
        run_4 = f'{funcionario.rg} '
        run_5 = 'inscrito no CPF/MF sob nº '
        # run_6 = '{}.{}.{}-{}, '.format(funcionario.cpf[:3], funcionario.cpf[3:6], funcionario.cpf[6:9], funcionario.cpf[9:])
        run_6 = funcionario.cpf
        run_7 = ' doravante denominado (a) '
        run_8 = '"CONTRATADO(A)" '
        run_9 = 'e, do outro lado, SOFTWARE AG BRASIL INFORMÁTICA E SERVIÇOS LTDA., ' \
                 'sociedade brasileira, estabelecida na Cidade de São Paulo, ' \
                 'Estado de São Paulo, na Avenida das Nações Unidas, 12.901, ' \
                 'Torre Norte, 33º andar, CEP 04578-000, Brooklin Novo, ' \
                 'inscrita no Cadastro Nacional das Pessoas Jurídicas (CNPJ/MF) sob no. 07.594.862/0001-39, ' \
                 'doravante denominada “SOFTWARE AG”;'

        par_1.add_run(run_1)
        par_1.add_run(run_2).bold = True
        par_1.add_run(run_3)
        par_1.add_run(run_4).bold = True
        par_1.add_run(run_5)
        par_1.add_run(run_6).bold = True
        par_1.add_run(run_7)
        par_1.add_run(run_8).bold = True
        par_1.add_run(run_9)

        par_2 = documento.add_paragraph()
        par_2.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        run_1 = 'Considerando que o(a) '
        run_2 = 'CONTRATADO(A) '
        run_3 = 'ocupa atualmente a função de '
        run_4 = f'{funcionario.cargo} '
        run_5 = 'pela '
        run_6 = 'SOFTWARE AG'
        run_7 = ', com data de admissão em '
        run_8 = '{}/{}/{}'.format(funcionario.admissao[8:], funcionario.admissao[5:7], funcionario.admissao[:4])
        run_9 = '.'

        par_2.add_run(run_1)
        par_2.add_run(run_2).bold = True
        par_2.add_run(run_3)
        par_2.add_run(run_4).bold = True
        par_2.add_run(run_5)
        par_2.add_run(run_6).bold = True
        par_2.add_run(run_7)
        par_2.add_run(run_8).bold = True
        par_2.add_run(run_9)

        par_3 = documento.add_paragraph()
        par_3.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        run_1 = 'Considerando que o (a) '
        run_2 = 'CONTRATADO(A) '
        run_3 = 'necessita utilizar aparelho celular para o exercício da sua função na '
        run_4 = 'SOFTWARE AG'
        run_5 = '.'

        par_3.add_run(run_1)
        par_3.add_run(run_2).bold = True
        par_3.add_run(run_3)
        par_3.add_run(run_4).bold = True
        par_3.add_run(run_5)

        par_4 = documento.add_paragraph()
        par_4.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        run_1 = 'Resolvem '.upper()
        run_2 = 'as partes firmar o presente instrumento que será regido pelas seguintes cláusulas e condições:'

        par_4.add_run(run_1)
        par_4.add_run(run_2)

        par_5 = documento.add_paragraph('', style='List Number')
        par_5.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        par_5.add_run('A SOFTWARE AG, neste ato, cede ao CONTRATADO o direito de uso de ')
        for idx, aparelho in enumerate(funcionario.aparelhos):
            if (idx > 0):
                par_5.add_run(', ')
            run_1 = '1 (um) aparelho telefônico marca '
            run_2 = f'{aparelho.marca}'
            run_3 = ', modelo '
            run_4 = f'{aparelho.modelo}'
            run_5 = ', IMEI '
            run_6 = f'{aparelho.imei}'
            run_7 = ', S/N '
            run_8 = '{}'.format(aparelho.numero_serie if aparelho.numero_serie else 'N/A')
            run_9 = ', doravante denominado "Celular", '
            run_10 = f'{aparelho.acessorios}'
            run_11 = ', 1 linha nº '
            run_12 = '({}) {}.{}-{}'.format(aparelho.linha.ddd, aparelho.linha.numero[:1], aparelho.linha.numero[1:5], aparelho.linha.numero[5:])
            par_5.add_run(run_1)
            par_5.add_run(run_2).bold = True
            par_5.add_run(run_3)
            par_5.add_run(run_4).bold = True
            par_5.add_run(run_3)
            par_5.add_run(run_5)
            par_5.add_run(run_6).bold = True
            par_5.add_run(run_7)
            par_5.add_run(run_8).bold = True
            par_5.add_run(run_9)
            par_5.add_run(run_10).bold = True
            par_5.add_run(run_11)
            par_5.add_run(run_12).bold = True
        par_5.add_run(' ')
        for idx, linha in enumerate(funcionario.linhas):
            if(idx > 0):
                par_5.add_run(', ')
            run_1 = '1 linha nº '
            run_2 = '({}) {}.{}-{}'.format(linha.ddd, linha.numero[:1], linha.numero[1:5], linha.numero[5:])
            par_5.add_run(run_1)
            par_5.add_run(run_2).bold = True
        
        run_1 = ', todos os itens de propriedade da SOFTWARE AG, '
        run_2 = 'exclusivamente para a prestação dos serviços pelo(a) CONTRATADO(A)'
        run_3 = ', não podendo ser dada outra finalidade, em qualquer hipótese.\t\n'
        par_5.add_run(run_1)
        par_5.add_run(run_2).bold = True
        par_5.add_run(run_3)

        par_6 = documento.add_paragraph('', style='List Number')
        par_6.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        run_1 = 'O(A) CONTRATADO(A) declara e reconhece que o direito de uso do Celular é '
        run_2 = 'exclusivo, pessoal e intransferível'
        run_3 = ', se comprometendo '
        run_4 = 'a não ceder, emprestar ou de qualquer forma permitir que terceiros utilizem o Celular'
        run_5 = '.\t\n'
        par_6.add_run(run_1)
        par_6.add_run(run_2).bold = True
        par_6.add_run(run_3)
        prun_4 = par_6.add_run(run_4)
        prun_4.bold = True
        prun_4.italic = True
        par_6.add_run(run_5)

        par_7 = documento.add_paragraph('', style='List Number')
        par_7.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        run_1 = 'A responsabilidade pelo '
        run_2 = 'uso apropriado '
        run_3 = 'do Celular disponibilizado pela SOFTWARE AG será '
        run_4 = 'única e exclusivamente do(a) CONTRATADO(A)'
        run_5 = '. Para fins desta cláusula, entende-se por '
        run_6 = '"uso apropriado"'
        run_7 = ' o uso em conformidade com a legislação em vigor e com as políticas internas da '
        run_8 = 'SOFTWARE AG'
        run_9 = ', inclusive Código de Ética e Conduta e suas políticas de segurança da informação. ' \
                'Em caso de dúvidas quanto ao '
        run_10 = ' ou políticas internas da '
        run_11 = ', o(a) '
        run_12 = 'CONTRATADO(A) '
        run_13 = 'deverá contatar ao departamento de TI, ao departamento de Proteção de Dados ou departamento de RH.\t\n'
        par_7.add_run(run_1)
        par_7.add_run(run_2).bold = True
        par_7.add_run(run_3)
        par_7.add_run(run_4).bold = True
        par_7.add_run(run_5)
        par_7.add_run(run_6).bold = True
        par_7.add_run(run_7)
        par_7.add_run(run_8).bold = True
        par_7.add_run(run_9)
        par_7.add_run(run_6).bold = True
        par_7.add_run(run_10)
        par_7.add_run(run_8).bold = True
        par_7.add_run(run_11)
        par_7.add_run(run_12).bold = True
        par_7.add_run(run_13)

        par_8 = documento.add_paragraph('', style='List Number')
        par_8.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        run_1 = 'O(A) '
        run_2 = 'CONTRATADO(A) '
        run_3 = 'têm ciência de que o uso do Celular, bem como os dados (inclusive pessoais) ' \
                'armazenados no Celular ou em nuvem relativos ao uso do Celular ' \
                'poderão ser acessados e/ou monitorados pela '
        run_4 = 'SOFTWARE AG'
        run_5 = ' para fins de (i) verificação do cumprimento deste termo; ' \
                '(ii) assistência técnica; (iii) eventual atualização de aplicativos; ' \
                '(iv) configuração/parametrização necessária para adequação às políticas de ' \
                'segurança da informação da '
        run_6 = '.\t\n'
        par_8.add_run(run_1)
        par_8.add_run(run_2).bold = True
        par_8.add_run(run_3)
        par_8.add_run(run_3)
        par_8.add_run(run_4).bold = True
        par_8.add_run(run_5)
        par_8.add_run(run_4).bold = True
        par_8.add_run(run_6)

        par_9 = documento.add_paragraph('', style='List Number')
        par_9.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        run_1 = 'O Celular e todos os acessórios descritos neste termo (Inclusive caixa) deverão ser devolvidos à '
        run_2 = 'SOFTWARE AG'
        run_3 = ' (ao departamento de TI) nas seguintes ocasiões:\t'
        par_9.add_run(run_1)
        par_9.add_run(run_2).bold = True
        par_9.add_run(run_3).add_break()
        prun_91 = par_9.add_run('\u2022 \u0009 Desligamento da SOFTWARE AG.\t')
        prun_91.bold = True
        prun_91.add_break()
        prun_92 = par_9.add_run('\u2022 \u0009 Transferência ou realocação do(a) CONTRATADO(A) a outra função.\t')
        prun_92.bold = True
        prun_92.add_break()
        prun_93 = par_9.add_run('\u2022 \u0009 Atualização do aparelho e/ou modelo.\t')
        prun_93.bold = True
        prun_93.add_break()
        prun_94 = par_9.add_run('\u2022 \u0009 Troca de aparelho por problema e/ou manutenção.\t')
        prun_94.bold = True
        prun_94.add_break()
        prun_95 = par_9.add_run('\u2022 \u0009 Solicitação da SOFTWARE AG, a qualquer tempo.\t\n')
        prun_95.bold = True
        
        par_10 = documento.add_paragraph('', style='List Number')
        par_10.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        run_1 = 'Caso o Celular e seus acessórios não sejam devolvidos corretamente, ' \
                'os respectivos custos serão de responsabilidade do(a) CONTRATADO(A), ' \
                'podendo referidos custos ser descontados de quaisquer verbas que a SOFTWARE AG ' \
                'tenha a haver ao(à) CONTRATADO(A), com exceção dos casos de perda, ' \
                'roubo ou furto, mediante a apresentação do boletim de ocorrência '
        run_2 = 'contendo o IMEI do Celular'
        run_3 = '.\t\n'
        par_10.add_run(run_1).bold = True
        prun_101 = par_10.add_run(run_2)
        prun_101.bold = True
        prun_101.underline = True
        par_10.add_run(run_3)

        par_11 = documento.add_paragraph('', style='List Number')
        par_11.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        run_1 = 'As partes elegem o foro da Comarca de São Paulo para dirimir quaisquer questões ' \
                'oriundas do presente instrumento. E por estarem assim justas e acordadas, ' \
                'as partes firmam o presente Contrato em 2 (duas) vias de igual teor. \t\n'
        par_11.add_run(run_1)

        data_atual = datetime.today().strftime('%Y-%m-%d')
        par_12 = documento.add_paragraph()
        par_12.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        run_1 = 'São Paulo, '
        run_2 = '{}'.format(data_atual[8:])
        run_3 = ' de '
        run_4 = ''
        if data_atual[5:7] == '01':
            run_4 = 'Janeiro'
        elif data_atual[5:7] == '02':
            run_4 = 'Fevereiro'
        elif data_atual[5:7] == '03':
            run_4 = 'Março'
        elif data_atual[5:7] == '04':
            run_4 = 'Abril'
        elif data_atual[5:7] == '05':
            run_4 = 'Maio'
        elif data_atual[5:7] == '06':
            run_4 = 'Junho'
        elif data_atual[5:7] == '07':
            run_4 = 'Julho'
        elif data_atual[5:7] == '08':
            run_4 = 'Agosto'
        elif data_atual[5:7] == '09':
            run_4 = 'Setembro'
        elif data_atual[5:7] == '10':
            run_4 = 'Outubro'
        elif data_atual[5:7] == '11':
            run_4 = 'Novembro'
        elif data_atual[5:7] == '12':
            run_4 = 'Dezembro'
        run_5 = ' de '
        run_6 = '{}'.format(data_atual[:4])
        par_12.add_run(run_1)
        par_12.add_run(run_2).bold = True
        par_12.add_run(run_3)
        par_12.add_run(run_4).bold = True
        par_12.add_run(run_5)
        par_12.add_run(run_6).bold = True

        documento.add_paragraph()
        # documento.add_paragraph()

        tbl_1 = documento.add_table(rows=2, cols=2)
        tbl_1.alignment = WD_TABLE_ALIGNMENT.CENTER
        tbl_1.autofit = True
        # tbl_1.rows[0].cells[0].text = 'Teste'
        tbl_1_ln_1 = tbl_1.rows[0].cells
        tbl_1_ln_1[0].text = '___________________________________________________'
        tbl_1_ln_1[1].text = '___________________________________________________'
        tbl_1_ln_2 = tbl_1.rows[1].cells
        tbl_1_ln_2[0].text = '{} {}'.format(funcionario.nome, funcionario.sobrenome).upper()
        tbl_1_ln_2[1].text = 'SOFTWARE AG BRASIL INFORMÁTICA E SERVIÇOS LTDA'

        # Salvando o Documento
        doc_path = os.path.join(app_config.UPLOAD_DOCS, f'termo_resp_{funcionario.id}.docx')
        documento.save(doc_path)

        # Abrindo o Documento para gravação local
        os.system(f'start {doc_path}')

